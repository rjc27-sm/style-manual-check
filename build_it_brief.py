"""Generate Style_Manual_Check_IT_Brief.docx using python-docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = "C:/Projects/style-manual-check/Style_Manual_Check_IT_Brief.docx"

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x38, 0x64)   # heading colour
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
GREY_FILL  = "D9E2F3"   # light blue-grey for table headers
WHITE_FILL = "FFFFFF"

def set_cell_bg(cell, hex_colour):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_colour)
    tcPr.append(shd)

def bold_cell(cell, text, size=10, colour=None):
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if colour:
        run.font.color.rgb = colour

def plain_cell(cell, text, size=10):
    p = cell.paragraphs[0]
    p.clear()
    p.add_run(text).font.size = Pt(size)

def h1(text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.color.rgb = MID_BLUE
        run.font.size = Pt(13)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(8)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix + ": ")
        run.bold = True
    p.add_run(text)
    return p

def spacer():
    p = doc.add_paragraph("")
    p.paragraph_format.space_after = Pt(2)
    return p

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Style Manual Check", 0)
for run in title.runs:
    run.font.color.rgb = DARK_BLUE
    run.font.size = Pt(22)

sub = doc.add_paragraph("IT briefing document")
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
for run in sub.runs:
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    run.font.size = Pt(12)
    run.italic = True
sub.paragraph_format.space_after = Pt(4)

meta = doc.add_paragraph("Prepared by Jen Robertson, Australian Institute of Health and Welfare (AIHW)  |  March 2026")
for run in meta.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
meta.paragraph_format.space_after = Pt(14)

doc.add_paragraph("")

# ── 1. Executive summary ──────────────────────────────────────────────────────
h1("1. Executive summary")
body(
    "Style Manual Check is an Australian Government tool that automatically checks written "
    "documents against the Australian Government Style Manual. It flags style issues — such "
    "as incorrect spelling, punctuation, date formatting, heading conventions, and the use of "
    "Latin abbreviations — and offers one-click fixes, each linked directly to the relevant "
    "Style Manual guidance. The tool has been developed by Jen Robertson at the Australian "
    "Institute of Health and Welfare (AIHW) and exists in two forms: a browser-based checker "
    "available publicly on GitHub Pages, and a Microsoft Word add-in that integrates directly "
    "into the document editing environment. The Word add-in is the intended long-term product; "
    "this document is provided to support the IT security and deployment review required before "
    "it can be made available to APS agency staff."
)

# ── 2. Purpose and goals ──────────────────────────────────────────────────────
h1("2. Purpose and goals")

h2("The problem")
body(
    "Australian Government staff produce large volumes of written material — policy documents, "
    "reports, ministerial briefs, web content — and compliance with the Style Manual is "
    "inconsistent. Manual checking is slow, dependent on individual knowledge, and easy to "
    "overlook under deadline pressure. Style errors in published government documents reduce "
    "professionalism, create accessibility issues (for example, headings that are bold text "
    "rather than properly styled are invisible to screen readers and navigation tools), and "
    "can cause confusion or embarrassment."
)

h2("The goal")
body(
    "Style Manual Check provides automated, real-time style checking embedded directly in "
    "the tools staff already use. It does not replace editorial judgement — it flags potential "
    "issues and lets the writer decide. Every suggestion links to the Style Manual so staff "
    "learn the rationale, not just the rule."
)

h2("Who benefits")
bullet("Writers and editors across APS agencies who produce official publications")
bullet("Communications teams responsible for consistency across large document suites")
bullet("New staff who are less familiar with the Style Manual")
bullet("Agencies aiming to improve accessibility compliance in their documents")
spacer()

h2("Ownership and maintenance")
body(
    "Proposed long-term owner: Australian Public Service Commission (APSC). "
    "Maintenance is estimated at approximately 1–2 days per month (rule updates, "
    "bug fixes, Style Manual changes). The source code is openly licensed (CC BY-NC 4.0) "
    "and hosted on GitHub."
)

# ── 3. What it checks ─────────────────────────────────────────────────────────
h1("3. What it checks")
body(
    "The tool contains 73 rules across 9 categories, plus a dictionary of over 500 "
    "US-to-Australian English spelling mappings (e.g. organize → organise, color → colour, "
    "center → centre). Rules are based on plain text analysis — there is no artificial "
    "intelligence or machine learning involved, and no calls to any external API or service."
)

# Rules table
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Set column widths
for i, width in enumerate([Cm(4.5), Cm(1.8), Cm(9.5)]):
    for cell in tbl.columns[i].cells:
        cell.width = width

hdr = tbl.rows[0].cells
bold_cell(hdr[0], "Category", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(hdr[1], "Rules", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(hdr[2], "Examples", colour=RGBColor(0xFF,0xFF,0xFF))
for c in hdr:
    set_cell_bg(c, "1F3864")

rules_data = [
    ("Abbreviations", "11",
     "Use 'for example' not e.g.; 'that is' not i.e.; 'and so on' not etc.; "
     "correct full stop usage in abbreviations and units"),
    ("Dates and time", "11",
     "US date format (January 15 → 15 January); ambiguous numeric dates; "
     "decade apostrophes (1980's → 1980s); time zone position; 12 am/pm ambiguity"),
    ("Government terms", "6",
     "Capitalisation of 'Australian Government'; minister/secretary preposition "
     "usage; references to generic departments and agencies"),
    ("Headings", "4",
     "Title case headings (should be sentence case); headings ending in full stops; "
     "all-caps headings; headings over 8 words"),
    ("Lists", "6",
     "Semicolons, commas, 'and'/'or' at end of list items; 'etc.' in lists; "
     "inconsistent capitalisation or punctuation across list items"),
    ("Numbers and measurements", "11",
     "Spell out zero and one; use numerals for 2 and above; "
     "ordinal words (1st → first for 1–9); per cent vs %; imperial units"),
    ("Punctuation", "11",
     "Em dashes (should be spaced en dashes); double quotation marks (should be single); "
     "double spaces; ampersands in body text; capital after colon"),
    ("Readability", "3",
     "Long sentences (over 35 words); watch words (words the Style Manual flags for "
     "consideration); wordy phrases (e.g. 'in order to' → 'to')"),
    ("Spelling", "9",
     "-ize → -ise; -or → -our; -er → -re; -ense → -ence; doubled consonants "
     "(traveled → travelled); gray → grey; 500+ US-to-AU mappings"),
]

for cat, count, examples in rules_data:
    row = tbl.add_row().cells
    plain_cell(row[0], cat)
    plain_cell(row[1], count)
    plain_cell(row[2], examples)
    # Alternate row shading
    if rules_data.index((cat, count, examples)) % 2 == 0:
        for c in row:
            set_cell_bg(c, "EEF3FA")

spacer()
body(
    "Each rule carries a direct link to the relevant section of the Style Manual "
    "(stylemanual.gov.au), displayed alongside the suggestion in the tool's interface."
)

# ── 4. Two versions ───────────────────────────────────────────────────────────
h1("4. The two versions")
body(
    "Style Manual Check is available in two forms. The browser version provides "
    "immediate access without installation; the Word add-in is the intended operational "
    "product for APS agencies."
)

# Comparison table
comp = doc.add_table(rows=1, cols=3)
comp.style = "Table Grid"
comp.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, width in enumerate([Cm(5.0), Cm(5.5), Cm(5.5)]):
    for cell in comp.columns[i].cells:
        cell.width = width

ch = comp.rows[0].cells
bold_cell(ch[0], "Feature", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(ch[1], "Browser version", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(ch[2], "Word add-in", colour=RGBColor(0xFF,0xFF,0xFF))
for c in ch:
    set_cell_bg(c, "1F3864")

comp_data = [
    ("How it's accessed",
     "Visit public URL in any browser",
     "Installed into Microsoft Word"),
    ("Where text is processed",
     "User's browser (client-side JS)",
     "Within Word on the user's device"),
    ("Text leaves the network?",
     "Yes — copied to a public website",
     "No — never leaves Word/M365"),
    ("Checks document formatting\n(heading styles, list styles)",
     "No — plain text only",
     "Yes — uses Word's style information for greater accuracy"),
    ("Fixes apply to the real document",
     "No — scratch pad only",
     "Yes — directly edits the document"),
    ("Handles sensitive documents",
     "Not recommended",
     "Yes — fully appropriate"),
    ("Installation required",
     "None",
     "Add-in deployment via M365 admin"),
    ("Suitable for official use",
     "With caution (public site)",
     "Yes"),
]

for i, (feat, browser, word) in enumerate(comp_data):
    row = comp.add_row().cells
    plain_cell(row[0], feat)
    plain_cell(row[1], browser)
    plain_cell(row[2], word)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "EEF3FA")

spacer()

h2("Browser version")
body(
    "The browser version is a single HTML page hosted on GitHub Pages "
    "(rjc27-sm.github.io/style-manual-check). Users paste text into the left panel, "
    "click 'Scan document', and review suggestions on the right. It is useful for "
    "quick, informal checks on non-sensitive text — particularly for staff who do not "
    "yet have access to the Word add-in. All processing happens in the user's browser "
    "using client-side JavaScript; no data is transmitted to any server. However, the "
    "act of copying text out of a secure environment and pasting it into a public website "
    "is a data-handling risk that agencies should consider carefully (see Section 5)."
)

h2("Word add-in")
body(
    "The Word add-in appears as a 'Style check' button in the Home tab ribbon. Clicking "
    "it opens a task pane on the right side of the document showing all detected issues, "
    "grouped by category. For each issue the user can: accept the suggestion (applies the "
    "fix directly to the document), use a specific replacement from alternatives offered, "
    "ignore the issue for this instance, navigate to the location in the document, or "
    "fix all instances of that rule at once. Because the add-in communicates with Word "
    "through the Office.js API, it has access to document structure — it knows which "
    "paragraphs have heading styles applied, which are list items, and which are formatted "
    "as bold — enabling more accurate checking than is possible with plain text alone."
)

# ── 5. Security and data handling ─────────────────────────────────────────────
h1("5. Security and data handling")

h2("Word add-in: data never leaves the network")
body(
    "When a user clicks 'Style check' in Word, the add-in reads the document text using "
    "the Office.js API — a standard Microsoft interface for Office add-ins — and passes "
    "it to the rule engine running locally within the Word task pane. No text is sent to "
    "any external server. No network calls are made during a check. The add-in contains "
    "no analytics, telemetry, logging, or external dependencies that would cause data to "
    "leave the device. This means the tool is appropriate for use with sensitive documents "
    "including policy drafts, budget material, and ministerial correspondence."
)

h2("Browser version: data-handling considerations")
body(
    "The browser version processes text entirely in the user's browser — there is no "
    "server receiving or storing the pasted content. However, to use it, staff must copy "
    "text from their work environment and paste it into a public website. Even though that "
    "website does not transmit or store the text, the act of copying content outside the "
    "agency network may conflict with information security policies, particularly for "
    "OFFICIAL or sensitive material. Agencies should assess this against their own policies "
    "and communicate appropriate use guidance to staff. The browser version is best suited "
    "to non-sensitive, publicly releasable text only."
)

h2("Add-in permissions")
body(
    "The Word add-in requests a single permission level: ReadWriteDocument. This is the "
    "standard permission for task pane add-ins that need to both read document content and "
    "apply edits. It grants access only to the currently open document. It does not grant "
    "access to:"
)
for item in [
    "other files or documents",
    "the user's file system",
    "email or calendar data",
    "network resources or the internet",
    "any system outside the Word document sandbox",
]:
    bullet(item)
spacer()
body(
    "The ReadWriteDocument permission is a standard, well-understood permission in the "
    "Office Add-ins security model and is required by all add-ins that apply edits to "
    "a document."
)

h2("No backend infrastructure")
body(
    "Style Manual Check has no server-side component, no database, no API, and no cloud "
    "service. It is entirely static files (HTML, JavaScript, a manifest file). There is "
    "nothing to patch, no service to monitor, and no user data stored anywhere."
)

# ── 6. Technical architecture ─────────────────────────────────────────────────
h1("6. Technical architecture")

h2("Rule engine")
body(
    "The core rule engine is approximately 3,000 lines of vanilla JavaScript (no framework "
    "dependencies). It accepts document text and returns a list of issues with positions, "
    "suggestions, and autofix strings. It also accepts optional metadata from Word — the "
    "set of line numbers that have heading styles, list styles, or bold formatting — which "
    "it uses to improve accuracy. The engine is entirely deterministic: the same input "
    "always produces the same output. There is no AI, no machine learning, and no "
    "probabilistic component."
)

h2("Word add-in platform")
body(
    "The add-in is built on Microsoft's Office Add-ins platform using the Office.js API — "
    "the same platform used by all Microsoft-approved Word, Excel, and PowerPoint add-ins. "
    "It is packaged as a task pane add-in: a small web application (HTML + JavaScript) "
    "that runs in an embedded browser frame inside Word. The add-in communicates with Word "
    "exclusively through the Office.js API; it cannot access the host operating system "
    "directly."
)

h2("Manifest file")
body(
    "A manifest file (manifest.xml) declares the add-in to Word. It specifies:"
)
for item in [
    "the add-in's unique identifier (GUID)",
    "the hosting URL where the add-in's files are served",
    "the permissions requested (ReadWriteDocument)",
    "the ribbon button definition (label, icon, action)",
    "the supported Office host (Word)",
]:
    bullet(item)
spacer()
body(
    "The manifest is the only file that needs to be deployed to users' machines or "
    "registered in the M365 admin centre. The actual add-in files (HTML and JavaScript) "
    "are served from a hosting location over HTTPS."
)

h2("Hosting requirement")
body(
    "The add-in's HTML and JavaScript files must be served over HTTPS from a location "
    "accessible to users. Suitable options include:"
)
bullet("An internal agency web server with HTTPS", bold_prefix=None)
bullet("SharePoint Online within the agency's M365 tenancy")
bullet("Azure Static Web Apps within the agency's Azure tenant")
bullet("Any static file host with a valid HTTPS certificate")
spacer()
body(
    "The files are entirely static — there is no server-side processing, no runtime "
    "environment to maintain, and no database. Hosting requirements are minimal."
)

h2("Build toolchain")
body(
    "The source code is compiled using Webpack and Babel (standard JavaScript build tools). "
    "The output is a small set of static files: taskpane.html, taskpane.js, commands.js, "
    "and manifest.xml. IT staff only need to host the compiled output — the build process "
    "is a developer concern and does not need to run in the agency environment."
)

h2("File sizes")
body(
    "The compiled add-in is small. The rule engine JavaScript file is approximately 140 KB; "
    "the spelling dictionary is approximately 20 KB. Total add-in size (all files) is well "
    "under 500 KB."
)

# ── 7. Deployment options ─────────────────────────────────────────────────────
h1("7. Deployment options (Word add-in)")
body(
    "There are three methods for deploying the Word add-in, in increasing order of IT "
    "involvement and reach."
)

h2("Option 1: Sideloading (current — testing only)")
body(
    "Individual users load the manifest file themselves via Word → Insert → Add-ins → "
    "Upload My Add-in. No IT involvement is required. This method is suitable only for "
    "development and user testing; it is not appropriate for production deployment as it "
    "requires each user to manually load and update the add-in."
)

h2("Option 2: Centralised Deployment (recommended for pilot)")
body(
    "An M365 administrator deploys the add-in via the Microsoft 365 Admin Centre "
    "(admin.microsoft.com → Settings → Integrated apps). The add-in is pushed to targeted "
    "users, groups, or the entire organisation and appears automatically in Word — no "
    "action required from end users. This is Microsoft's recommended deployment method "
    "for organisational add-ins. Requirements:"
)
for item in [
    "M365 administrator access",
    "A hosted manifest URL (HTTPS)",
    "Microsoft 365 Apps for enterprise licences (standard APS provision)",
]:
    bullet(item)

h2("Option 3: AppSource or SharePoint App Catalog (broad release)")
body(
    "For APS-wide availability, the add-in can be published to Microsoft AppSource "
    "(Microsoft's commercial add-in marketplace) or to a SharePoint App Catalog. "
    "AppSource publication requires a Microsoft Partner account and a security review "
    "by Microsoft. A SharePoint App Catalog provides an agency-internal alternative "
    "without marketplace publication."
)

# ── 8. Roadmap ────────────────────────────────────────────────────────────────
h1("8. Current status and roadmap")
body("The project is currently in working-group testing at AIHW.")

# Roadmap table
rt = doc.add_table(rows=1, cols=3)
rt.style = "Table Grid"
rt.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, width in enumerate([Cm(3.5), Cm(2.5), Cm(10.0)]):
    for cell in rt.columns[i].cells:
        cell.width = width

rh = rt.rows[0].cells
bold_cell(rh[0], "Phase", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(rh[1], "Timing", colour=RGBColor(0xFF,0xFF,0xFF))
bold_cell(rh[2], "Description", colour=RGBColor(0xFF,0xFF,0xFF))
for c in rh:
    set_cell_bg(c, "1F3864")

roadmap = [
    ("Working group testing", "Q1 2026",
     "Small group of writers and editors testing the add-in via sideloading at AIHW. "
     "Gathering feedback on rule accuracy and usability."),
    ("IT review", "Q2 2026",
     "Security and deployment review by APSC and AIHW IT. This document supports "
     "that review."),
    ("Pilot agencies", "Q3 2026",
     "Centralised Deployment to one or two pilot agencies. Broader user testing "
     "with real documents in a production M365 environment."),
    ("Broad release", "Q4 2026",
     "APS-wide availability via AppSource or SharePoint App Catalog. "
     "Proposed ownership transferred to APSC."),
    ("PowerPoint add-in", "Post Q4 2026",
     "A PowerPoint version using the same rule engine is planned after the "
     "Word add-in is validated in production."),
]

for i, (phase, timing, desc) in enumerate(roadmap):
    row = rt.add_row().cells
    plain_cell(row[0], phase)
    plain_cell(row[1], timing)
    plain_cell(row[2], desc)
    if i % 2 == 0:
        for c in row:
            set_cell_bg(c, "EEF3FA")

# ── 9. IT requirements and open questions ─────────────────────────────────────
h1("9. What IT needs to consider")

h2("For the IT review")
for item in [
    ("Security review of the add-in codebase",
     "The source code is open and available at github.com/rjc27-sm/style-manual-check. "
     "The compiled add-in files can be provided separately for review."),
    ("Hosting location",
     "A decision is needed on where to host the add-in files (internal web server, "
     "SharePoint, or Azure Static Web Apps). The files are static and require only HTTPS."),
    ("Centralised Deployment availability",
     "Confirm that Centralised Deployment is enabled in the agency's M365 tenancy and "
     "that the relevant admin has access to the Microsoft 365 Admin Centre."),
    ("Network/firewall requirements",
     "The add-in itself makes no outbound network calls. However, Word's Add-ins "
     "framework may require access to Microsoft CDN endpoints "
     "(appsforoffice.microsoft.com, az158161.vo.msecnd.net) for the Office.js library. "
     "These are standard Microsoft endpoints and are typically already permitted."),
    ("Branding and iconography",
     "The current add-in uses placeholder icons. Determine whether agency-specific "
     "branding is required before deployment."),
    ("Update and maintenance process",
     "Establish a process for deploying rule updates. Updates are applied by replacing "
     "the hosted files — no reinstallation or re-deployment of the manifest is needed "
     "for minor updates. Changes to the manifest itself would require a re-deployment "
     "via the Admin Centre."),
]:
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(item[0] + ": ")
    run.bold = True
    p.add_run(item[1])

# ── 10. Licensing ──────────────────────────────────────────────────────────────
h1("10. Licensing and source code")
body(
    "Style Manual Check is released under the Creative Commons "
    "Attribution-NonCommercial 4.0 International licence (CC BY-NC 4.0). "
    "Any person or organisation may use, adapt, and redistribute the tool provided "
    "they give appropriate credit and do not use it for commercial purposes. "
    "APS agencies may deploy and adapt the tool freely under this licence."
)
bullet("Source code: github.com/rjc27-sm/style-manual-check")
bullet("Browser version: rjc27-sm.github.io/style-manual-check")
bullet("Style Manual: stylemanual.gov.au")
bullet("Office Add-ins platform documentation: learn.microsoft.com/en-us/office/dev/add-ins")

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
